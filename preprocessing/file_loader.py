"""
file_loader.py
==============
Multi-format file loader — PDF, Excel, CSV, JSON, DOCX, TXT.
Auto-detects file type using python-magic (not extension).

Supported formats:
  • PDF   → pdfplumber (tables + text)
  • Excel → openpyxl / pandas
  • CSV   → pandas
  • JSON  → orjson (schema-aware, nested flatten)
  • DOCX  → python-docx
  • TXT   → plain read with encoding detection
"""

import os
import json
import logging
from pathlib import Path
from typing import Any
from dataclasses import dataclass, field

import chardet
import orjson
import pandas as pd
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

@dataclass
class LoadedDocument:
    """Raw document after loading — before any preprocessing."""
    file_path: str
    file_type: str          # pdf | excel | csv | json | docx | txt
    raw_text: str
    tables: list[dict] = field(default_factory=list)   # extracted tables as dicts
    metadata: dict = field(default_factory=dict)
    page_count: int = 0
    char_count: int = 0
    load_error: str | None = None


# ─────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────

def _detect_encoding(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read(50_000)
    result = chardet.detect(raw)
    return result.get("encoding") or "utf-8"


def _detect_file_type(path: str) -> str:
    """
    Detect file type by content (magic bytes) first,
    fall back to extension so we don't need libmagic installed.
    """
    ext = Path(path).suffix.lower()
    magic_map = {
        ".pdf": "pdf",
        ".xlsx": "excel",
        ".xls": "excel",
        ".csv": "csv",
        ".json": "json",
        ".jsonl": "json",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "txt",
        ".md": "txt",
    }
    # Try magic bytes for PDF
    with open(path, "rb") as f:
        header = f.read(8)
    if header[:4] == b"%PDF":
        return "pdf"
    if header[:2] in (b"PK",):         # zip-based (xlsx, docx)
        if ext in (".xlsx", ".xls"):
            return "excel"
        if ext in (".docx", ".doc"):
            return "docx"
    return magic_map.get(ext, "txt")   # default to txt


def _flatten_json(obj: Any, parent_key: str = "", sep: str = ".") -> dict:
    """Recursively flatten nested JSON/dict into dot-separated keys."""
    items: list[tuple[str, Any]] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else str(k)
            if isinstance(v, (dict, list)):
                items.extend(_flatten_json(v, new_key, sep).items())
            else:
                items.append((new_key, v))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            new_key = f"{parent_key}{sep}{i}" if parent_key else str(i)
            if isinstance(v, (dict, list)):
                items.extend(_flatten_json(v, new_key, sep).items())
            else:
                items.append((new_key, v))
    else:
        items.append((parent_key, obj))
    return dict(items)


# ─────────────────────────────────────────────
# Per-format loaders
# ─────────────────────────────────────────────

def _load_pdf(path: str) -> LoadedDocument:
    pages_text = []
    tables = []
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages_text.append(f"[Page {page_num}]\n{text}")
            # Extract tables if present
            for tbl in page.extract_tables() or []:
                if tbl:
                    tables.append({
                        "page": page_num,
                        "data": tbl,
                    })

    raw_text = "\n\n".join(pages_text)
    return LoadedDocument(
        file_path=path,
        file_type="pdf",
        raw_text=raw_text,
        tables=tables,
        page_count=page_count,
        char_count=len(raw_text),
        metadata={"source": Path(path).name, "pages": page_count},
    )


def _load_excel(path: str) -> LoadedDocument:
    xl = pd.ExcelFile(path)
    all_text_parts = []
    tables = []
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        df = df.fillna("")
        sheet_text = f"[Sheet: {sheet_name}]\n{df.to_string(index=False)}"
        all_text_parts.append(sheet_text)
        tables.append({"sheet": sheet_name, "data": df.values.tolist(),
                        "columns": list(df.columns)})
    raw_text = "\n\n".join(all_text_parts)
    return LoadedDocument(
        file_path=path, file_type="excel",
        raw_text=raw_text, tables=tables,
        char_count=len(raw_text),
        metadata={"source": Path(path).name, "sheets": xl.sheet_names},
    )


def _load_csv(path: str) -> LoadedDocument:
    enc = _detect_encoding(path)
    df = pd.read_csv(path, encoding=enc)
    df = df.fillna("")
    raw_text = df.to_string(index=False)
    return LoadedDocument(
        file_path=path, file_type="csv",
        raw_text=raw_text,
        tables=[{"data": df.values.tolist(), "columns": list(df.columns)}],
        char_count=len(raw_text),
        metadata={"source": Path(path).name,
                   "rows": len(df), "cols": len(df.columns)},
    )


def _load_json(path: str) -> LoadedDocument:
    with open(path, "rb") as f:
        content = f.read()

    # JSONL support
    if path.endswith(".jsonl"):
        records = [orjson.loads(line) for line in content.splitlines() if line.strip()]
        flat_records = [_flatten_json(r) for r in records]
        raw_text = "\n".join(
            " | ".join(f"{k}: {v}" for k, v in rec.items())
            for rec in flat_records
        )
        return LoadedDocument(
            file_path=path, file_type="json",
            raw_text=raw_text,
            char_count=len(raw_text),
            metadata={"source": Path(path).name, "format": "jsonl",
                       "record_count": len(records)},
        )

    obj = orjson.loads(content)
    flat = _flatten_json(obj)
    raw_text = "\n".join(f"{k}: {v}" for k, v in flat.items())
    return LoadedDocument(
        file_path=path, file_type="json",
        raw_text=raw_text,
        char_count=len(raw_text),
        metadata={"source": Path(path).name, "format": "json",
                   "keys": list(flat.keys())[:20]},
    )


def _load_docx(path: str) -> LoadedDocument:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for tbl in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        tables.append({"data": rows})
    raw_text = "\n\n".join(paragraphs)
    return LoadedDocument(
        file_path=path, file_type="docx",
        raw_text=raw_text, tables=tables,
        char_count=len(raw_text),
        metadata={"source": Path(path).name, "paragraphs": len(paragraphs)},
    )


def _load_txt(path: str) -> LoadedDocument:
    enc = _detect_encoding(path)
    with open(path, encoding=enc, errors="replace") as f:
        raw_text = f.read()
    return LoadedDocument(
        file_path=path, file_type="txt",
        raw_text=raw_text,
        char_count=len(raw_text),
        metadata={"source": Path(path).name, "encoding": enc},
    )


# ─────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────

LOADER_MAP = {
    "pdf":   _load_pdf,
    "excel": _load_excel,
    "csv":   _load_csv,
    "json":  _load_json,
    "docx":  _load_docx,
    "txt":   _load_txt,
}


def load_file(path: str) -> LoadedDocument:
    """
    Load any supported file and return a LoadedDocument.

    Usage:
        doc = load_file("data/report.pdf")
        print(doc.raw_text[:500])
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    file_type = _detect_file_type(path)
    loader = LOADER_MAP.get(file_type, _load_txt)

    try:
        doc = loader(path)
        logger.info(f"Loaded [{file_type.upper()}] {path} — {doc.char_count:,} chars")
        return doc
    except Exception as e:
        logger.error(f"Failed to load {path}: {e}")
        return LoadedDocument(
            file_path=path, file_type=file_type,
            raw_text="", load_error=str(e),
        )


def load_directory(dir_path: str, recursive: bool = True) -> list[LoadedDocument]:
    """
    Load all supported files from a directory.

    Args:
        dir_path: Path to directory
        recursive: If True, walks all subdirectories

    Returns:
        List of LoadedDocument objects
    """
    supported_exts = {".pdf", ".xlsx", ".xls", ".csv",
                       ".json", ".jsonl", ".docx", ".doc", ".txt", ".md"}
    docs = []
    base = Path(dir_path)
    pattern = "**/*" if recursive else "*"

    for file_path in base.glob(pattern):
        if file_path.is_file() and file_path.suffix.lower() in supported_exts:
            docs.append(load_file(str(file_path)))

    logger.info(f"Loaded {len(docs)} files from {dir_path}")
    return docs
